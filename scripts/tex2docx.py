"""Render paper/main.tex as a Word document for reading and comment.

The markdown mirrors are frozen at the pre-reduction draft, so they no longer show where anything
lives. This reads the tex itself, which is now the only canonical source.

It is a reading copy, not a typesetting engine: math is shown as its source with the delimiters
stripped, floats appear where they are declared rather than where LaTeX would place them, and
citations are rendered as author keys. Section numbers are computed, so the numbering in the
document matches what a compile would produce.

    python scripts/tex2docx.py paper/main.tex out.docx
"""
from __future__ import annotations

import io
import re
import sys

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

B = chr(92)
BS = B + B                     # a literal backslash inside a regex

DROP_LINE = re.compile(r"^" + BS + r"(?:documentclass|usepackage|def|newcommand|title|author|maketitle"
                       r"|begin\{document\}|end\{document\}|bibliography|bibliographystyle|appendix"
                       r"|vspace|centering|small|scriptsize|footnotesize|setlength|toprule|midrule"
                       r"|bottomrule|cmidrule|resizebox|input|includegraphics|label)\b")
INLINE = [
    (re.compile(r"" + BS + r"(?:textbf|textit|emph|texttt|mathbf|boldsymbol|mathrm|text)\{([^{}]*)\}"), r"\1"),
    (re.compile(r"" + BS + r"cite[a-z]*\*?(?:\[[^\]]*\])*\{([^}]*)\}"), r"[\1]"),
    (re.compile(r"" + BS + r"(?:ref|autoref)\{([^}]*)\}"), r"[\1]"),
    (re.compile(r"" + BS + r"(?:S|yes|dagger|ast|approx|times|le|ge|to|cup|in|rho|delta|lambda|hat|mid)\b"), ""),
    (re.compile(r"[$]"), ""),
    (re.compile(r"" + BS + r"%"), "%"),
    (re.compile(r"" + BS + r"&"), "&"),
    (re.compile(r"" + BS + r"_"), "_"),
    (re.compile(r"---"), "—"),
    (re.compile(r"``|''"), '"'),
    (re.compile(r"~"), " "),
    (re.compile(r"" + BS + r"[a-zA-Z]+"), ""),
    (re.compile(r"\{|\}"), ""),
    (re.compile(r"  +"), " "),
]


def clean(t):
    for pat, rep in INLINE:
        t = pat.sub(rep, t)
    return t.strip()


def main(src, dst):
    lines = io.open(src, encoding="utf-8").read().split(chr(10))
    doc = Document()
    doc.styles["Normal"].font.size = Pt(10.5)
    sec = [0, 0]
    in_appendix = False
    i, buf, table = 0, [], None

    def flush():
        if buf:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.add_run(clean(" ".join(buf)))
            buf.clear()

    while i < len(lines):
        ln = lines[i].strip()
        if ln.startswith(B + "appendix"):
            in_appendix, sec = True, [0, 0]
        m = re.match(r"^" + BS + r"(section|subsection)\*?\{(.*?)\}", ln)
        if m:
            flush()
            if m.group(1) == "section":
                sec[0] += 1; sec[1] = 0
                tag = (chr(64 + sec[0]) if in_appendix else str(sec[0]))
                doc.add_heading(f"{tag}. {clean(m.group(2))}", level=1)
            else:
                sec[1] += 1
                tag = (chr(64 + sec[0]) if in_appendix else str(sec[0]))
                doc.add_heading(f"{tag}.{sec[1]} {clean(m.group(2))}", level=2)
            i += 1; continue
        if ln.startswith(B + "paragraph{"):
            flush()
            buf.append(re.sub(r"^" + BS + r"paragraph\{(.*?)\}", r"\1.", ln))
            i += 1; continue
        if ln.startswith(B + "caption{"):
            flush()
            cap = doc.add_paragraph()
            cap.add_run(clean(ln[len(B + "caption{"):].rstrip("}"))).italic = True
            i += 1; continue
        if ln.startswith(B + "begin{tabular}"):
            flush()
            rows, i = [], i + 1
            while i < len(lines) and not lines[i].strip().startswith(B + "end{tabular}"):
                r = lines[i].strip()
                if r and not DROP_LINE.match(r):
                    cells = [clean(c) for c in re.split(r"(?<!" + BS + r")&", r.rstrip(B + B))]
                    if any(cells): rows.append(cells)
                i += 1
            if rows:
                w = max(len(r) for r in rows)
                t = doc.add_table(rows=0, cols=w); t.style = "Table Grid"
                for r in rows:
                    cs = t.add_row().cells
                    for k, c in enumerate(r[:w]):
                        cs[k].text = c
                        for run in cs[k].paragraphs[0].runs: run.font.size = Pt(8)
                doc.add_paragraph()
            i += 1; continue
        if not ln or ln.startswith("%") or DROP_LINE.match(ln) or ln.startswith(B + "begin") or ln.startswith(B + "end"):
            flush(); i += 1; continue
        buf.append(ln)
        i += 1
    flush()
    doc.save(dst)
    print("wrote", dst)


if __name__ == "__main__":
    main(sys.argv[1], sys.argv[2])
