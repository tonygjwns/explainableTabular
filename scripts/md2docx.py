"""Render a manuscript markdown file as a Word document for reading and comment.

This is a reading copy, not a deliverable: the submission artifact is paper/main.tex. It exists
because reviewing 1,300 lines of markdown in an editor is worse than reviewing a paginated document
with real tables, and because comments come back more easily on a .docx.

Korean text needs an East Asian font set explicitly, or Word substitutes per-run and the result
looks broken.

    python scripts/md2docx.py PAPER_DRAFT_V5_KO.md out.docx
"""
from __future__ import annotations

import io
import re
import sys

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt

BOLD = re.compile(r"\*\*([^*]+)\*\*")
ITAL = re.compile(r"(?<!\*)\*([^*\n]+)\*(?!\*)")
CODE = re.compile(r"`([^`]+)`")


def set_fonts(doc, latin="Calibri", east="Malgun Gothic"):
    style = doc.styles["Normal"]
    style.font.name = latin
    style.font.size = Pt(10.5)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), east)


def add_runs(par, text):
    """Emit bold, italic and inline-code spans as runs; everything else plain.

    Italics matter as much as bold: leaving single asterisks unhandled printed 229 literal stars
    into the first Korean reading copy.
    """
    text = CODE.sub(r"\1", text)
    text = text.replace(chr(92) + "*", chr(1))          # escaped star: not markup (D\* etc.)
    marks = [(m.start(), m.end(), m.group(1), "b") for m in BOLD.finditer(text)]
    for m in ITAL.finditer(text):
        if not any(a <= m.start() < b for a, b, _, _ in marks):
            marks.append((m.start(), m.end(), m.group(1), "i"))
    marks.sort()
    pos = 0
    for a, b, inner, kind in marks:
        if a > pos:
            par.add_run(text[pos:a].replace(chr(1), chr(42)))
        run = par.add_run(inner.replace(chr(1), chr(42)))
        if kind == "b":
            run.bold = True
        else:
            run.italic = True
        pos = b
    if pos < len(text):
        par.add_run(text[pos:].replace(chr(1), chr(42)))


def main(src, dst):
    lines = io.open(src, encoding="utf-8").read().split("\n")
    doc = Document()
    set_fonts(doc)
    i, buf = 0, []

    def flush():
        if buf:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            add_runs(p, " ".join(buf).strip())
            buf.clear()

    while i < len(lines):
        ln = lines[i]
        s = ln.strip()
        if not s:
            flush(); i += 1; continue
        if s.startswith(">"):        # internal build note, not part of the manuscript
            flush()
            while i < len(lines) and lines[i].strip().startswith(">"):
                i += 1
            continue
        if s.startswith("#"):
            flush()
            lvl = len(s) - len(s.lstrip("#"))
            doc.add_heading(re.sub(r"^#+\s*", "", s), level=min(lvl, 4))
            i += 1; continue
        if s.startswith("---") and set(s) <= {"-"}:
            flush(); doc.add_page_break(); i += 1; continue
        if s.startswith("|") and s.endswith("|"):   # a wrapped $|x|$ is not a table row
            flush()
            rows = []
            while i < len(lines) and lines[i].strip().startswith("|") and lines[i].strip().endswith("|"):
                cells = [c.strip() for c in lines[i].strip().strip("|").split("|")]
                if not all(set(c) <= {"-", ":", " "} for c in cells if c):
                    rows.append(cells)
                i += 1
            if rows:
                w = max(len(r) for r in rows)
                t = doc.add_table(rows=0, cols=w)
                t.style = "Table Grid"
                for r in rows:
                    cells = t.add_row().cells
                    for k, c in enumerate(r[:w]):
                        cells[k].text = ""
                        add_runs(cells[k].paragraphs[0], c)
                        for run in cells[k].paragraphs[0].runs:
                            run.font.size = Pt(8.5)
                doc.add_paragraph()
            continue
        buf.append(s)
        i += 1
    flush()
    doc.save(dst)
    print("wrote", dst)


if __name__ == "__main__":
    main(sys.argv[1], sys.argv[2])
